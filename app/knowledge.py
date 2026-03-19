from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import numpy as np
from docx import Document
from openai import OpenAI
from pypdf import PdfReader


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, max_chars: int = 1600, overlap: int = 250) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        candidate = f"{current}\n\n{para}".strip() if current else para
        if len(candidate) <= max_chars:
            current = candidate
            continue

        if current:
            chunks.append(current)

        if len(para) <= max_chars:
            current = para
            continue

        start = 0
        while start < len(para):
            end = start + max_chars
            chunks.append(para[start:end])
            if end >= len(para):
                break
            start = max(0, end - overlap)
        current = ""

    if current:
        chunks.append(current)

    return chunks


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages)


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def read_supported_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    return ""


def iter_supported_files(root: Path, extensions: set[str]) -> Iterable[Path]:
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if path.name.startswith(".") or path.parent.name.startswith("."):
            continue
        if path.suffix.lower() in extensions:
            yield path


@dataclass
class IndexedChunk:
    id: str
    source: str
    text: str
    embedding: list[float]


class KnowledgeBase:
    def __init__(self, index_path: str, client: OpenAI, embedding_model: str):
        self.index_path = Path(index_path)
        self.client = client
        self.embedding_model = embedding_model
        self.chunks: list[IndexedChunk] = []
        self.last_build_manifest: dict = {}
        if self.index_path.exists():
            self.load()

    def load(self) -> None:
        payload = json.loads(self.index_path.read_text(encoding="utf-8"))
        self.chunks = [IndexedChunk(**item) for item in payload.get("chunks", [])]
        self.last_build_manifest = payload.get("manifest", {})

    def save(self) -> None:
        self.index_path.parent.mkdir(parents=True, exist_ok=True)
        payload = {
            "chunks": [chunk.__dict__ for chunk in self.chunks],
            "manifest": self.last_build_manifest,
        }
        self.index_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    def compute_manifest(self, root: Path, extensions: set[str]) -> dict:
        files = []
        for path in sorted(iter_supported_files(root, extensions)):
            stat = path.stat()
            files.append(
                {
                    "path": str(path.relative_to(root.parent)),
                    "size": stat.st_size,
                    "mtime_ns": stat.st_mtime_ns,
                }
            )
        return {"root": str(root), "extensions": sorted(extensions), "files": files}

    def needs_rebuild(self, directory: str, extensions: set[str]) -> bool:
        root = Path(directory)
        if not self.index_path.exists():
            return True
        return self.compute_manifest(root, extensions) != self.last_build_manifest

    def build_from_directory(self, directory: str, extensions: set[str] | None = None) -> dict:
        root = Path(directory)
        if not root.exists():
            raise FileNotFoundError(f"Knowledge directory not found: {root}")

        extensions = extensions or {".txt", ".md", ".pdf", ".docx"}
        all_chunks: list[tuple[str, str, str]] = []
        file_count = 0
        skipped_files: list[str] = []

        for path in sorted(iter_supported_files(root, extensions)):
            file_count += 1
            try:
                raw = read_supported_file(path)
            except Exception as exc:
                skipped_files.append(f"{path.name}: {exc}")
                continue

            chunks = chunk_text(raw)
            if not chunks:
                skipped_files.append(f"{path.name}: no extractable text found")
                continue

            for idx, chunk in enumerate(chunks):
                chunk_id = hashlib.sha1(f"{path}:{idx}:{chunk}".encode("utf-8")).hexdigest()
                all_chunks.append((chunk_id, str(path.relative_to(root.parent)), chunk))

        self.chunks = []
        batch_size = 64
        for i in range(0, len(all_chunks), batch_size):
            batch = all_chunks[i : i + batch_size]
            embeddings = self.client.embeddings.create(
                model=self.embedding_model,
                input=[item[2] for item in batch],
            )
            for item, emb in zip(batch, embeddings.data):
                self.chunks.append(
                    IndexedChunk(
                        id=item[0],
                        source=item[1],
                        text=item[2],
                        embedding=emb.embedding,
                    )
                )

        self.last_build_manifest = self.compute_manifest(root, extensions)
        self.save()
        return {
            "files_indexed": file_count,
            "chunks_indexed": len(self.chunks),
            "index_path": str(self.index_path),
            "skipped_files": skipped_files,
        }

    def ensure_index(self, directory: str, extensions: set[str]) -> dict:
        if self.needs_rebuild(directory, extensions):
            result = self.build_from_directory(directory, extensions)
            result["rebuilt"] = True
            return result
        return {
            "files_indexed": len(self.last_build_manifest.get("files", [])),
            "chunks_indexed": len(self.chunks),
            "index_path": str(self.index_path),
            "skipped_files": [],
            "rebuilt": False,
        }

    def search(self, query: str, top_k: int = 5) -> list[dict]:
        if not self.chunks:
            return []

        query_emb = self.client.embeddings.create(
            model=self.embedding_model,
            input=query,
        ).data[0].embedding

        q = np.array(query_emb)
        scored: list[dict] = []
        for chunk in self.chunks:
            v = np.array(chunk.embedding)
            denom = np.linalg.norm(q) * np.linalg.norm(v)
            score = float(np.dot(q, v) / denom) if denom else 0.0
            scored.append(
                {
                    "source": chunk.source,
                    "text": chunk.text,
                    "score": round(score, 4),
                }
            )

        scored.sort(key=lambda item: item["score"], reverse=True)
        return scored[:top_k]
